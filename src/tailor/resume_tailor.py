# Resume tailoring module
# Uses Claude API to identify keyword swaps, then python-docx to apply them

import json
import re
import shutil
from pathlib import Path

_JUNK_LABEL = re.compile(
    r"\b(experience|education|skills|projects?|certifications?|summary|objective|"
    r"references?|awards?|leadership|technical\s+capabilities?|selected\s+projects?|"
    r"additional\s+information?)\b",
    re.IGNORECASE,
)


def _is_junk_paragraph(text: str) -> bool:
    """True if paragraph consists entirely of section-label words and separators."""
    if not text:
        return True
    cleaned = _JUNK_LABEL.sub("", text)
    cleaned = re.sub(r"[\s/\-\u2013\u2014|•→>]+", "", cleaned)
    return cleaned == ""


from docx import Document
from docx.shared import Pt

PROMPT_PATH = Path(__file__).parent.parent.parent / "prompts" / "resume_tailor.md"


def _extract_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _strip_trailing_junk(doc: Document) -> None:
    """Remove empty/junk trailing paragraphs and eliminate blank-page spacing."""
    for para in reversed(doc.paragraphs):
        if _is_junk_paragraph(para.text.strip()):
            try:
                para._element.getparent().remove(para._element)
            except Exception:
                pass  # mandatory final paragraph can't be deleted; leave it
        else:
            # Zero out space_after on the last real paragraph so it can't
            # push a blank second page due to trailing whitespace/spacing.
            para.paragraph_format.space_after = Pt(0)
            break


def _apply_replacements(doc: Document, replacements: list[dict]) -> None:
    """Apply find/replace pairs to all runs in the document, preserving formatting."""
    targets = [(r["find"], r["replace"]) for r in replacements if r.get("find")]

    def replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                for find, replace in targets:
                    if find in run.text:
                        run.text = run.text.replace(find, replace)

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)


def tailor_resume(job: dict, resume_path: str, output_path: str, client) -> str:
    """
    Tailor the DOCX resume for a specific job by swapping keywords.
    Preserves all formatting — only changes specific words/phrases.

    Args:
        job: Job dict with description, title, company, match analysis
        resume_path: Path to the base resume.docx
        output_path: Where to save the tailored resume.docx
        client: Anthropic client instance

    Returns:
        Path to the saved tailored resume.docx
    """
    resume_text = _extract_text(resume_path)
    prompt_template = PROMPT_PATH.read_text()

    recommended = job.get("recommended_keywords", [])
    substitutions = {
        "resume_text": resume_text,
        "company": job.get("company", ""),
        "job_title": job.get("title", ""),
        "job_description": job.get("description", ""),
        "recommended_keywords": ", ".join(recommended) if recommended else "None",
    }
    prompt = prompt_template
    for key, value in substitutions.items():
        prompt = prompt.replace(f"{{{key}}}", value)

    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    text = message.content[0].text

    # Extract JSON array from response
    json_match = re.search(r"\[.*\]", text, re.DOTALL)
    if not json_match:
        raise ValueError(f"No JSON array found in Claude response:\n{text}")

    replacements = json.loads(json_match.group())

    # Copy base resume and apply swaps
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(resume_path, output_path)

    doc = Document(output_path)
    _apply_replacements(doc, replacements)
    _strip_trailing_junk(doc)
    doc.save(output_path)

    return output_path
