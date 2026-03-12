# Cover letter tailoring module
# Haiku receives the template with numbered paragraphs + resume + JD,
# then returns JSON specifying which paragraph indices to update.
# Paragraph-index replacement preserves all DOCX formatting (fonts, spacing, layout).

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

# Matches bracket placeholders that are clearly template instructions, e.g.:
#   [Company Name]  [specific mission, product, or initiative]  [Role Title]
# Excludes numeric indices like [0] which are not template placeholders.
_TEMPLATE_BRACKET = re.compile(r'\[[A-Za-z][^\]]{2,}\]')

PROMPT_PATH = Path(__file__).parent.parent.parent / "prompts" / "cover_letter_tailor.md"


def _numbered_paragraphs(doc: Document) -> str:
    """Return non-empty template paragraphs as '[index] text' lines for the prompt."""
    lines = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            lines.append(f"[{i}] {para.text}")
    return "\n".join(lines)


def _apply_paragraph_replacements(doc: Document, replacements: list[dict]) -> None:
    """Replace paragraph content by index; physically delete paragraphs with text=''."""
    # Snapshot references before any mutation (deletions shift the live list)
    paras = doc.paragraphs
    to_process = []
    for item in replacements:
        idx = item.get("index")
        new_text = item.get("text", "")
        if idx is None or not isinstance(idx, int) or idx >= len(paras):
            continue
        to_process.append((paras[idx], new_text))

    for para, new_text in to_process:
        if new_text == "":
            # Remove the paragraph element entirely (eliminates blank-line spacing)
            para._element.getparent().remove(para._element)
        else:
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.add_run(new_text)


def _remove_template_remnants(doc: Document) -> None:
    """Delete any paragraph still containing unfilled template bracket placeholders.

    Haiku sometimes writes the tailored content to the correct index but forgets
    to clear the adjacent original template paragraph. This catches what it misses.
    """
    for para in list(doc.paragraphs):
        if _TEMPLATE_BRACKET.search(para.text):
            try:
                para._element.getparent().remove(para._element)
            except Exception:
                pass


def _apply_formatting(doc: Document, font_pt: float = 11.0, margin_top: float = 0.5, margin_bottom: float = 0.5) -> None:
    """Set font size on all runs and tighten top/bottom page margins."""
    size = Pt(font_pt)
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = size
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = size
    for section in doc.sections:
        section.top_margin = Inches(margin_top)
        section.bottom_margin = Inches(margin_bottom)


def tailor_cover_letter(job: dict, cover_letter_path: str, output_path: str, client, resume_text: str = "") -> str:
    """
    Tailor the DOCX cover letter for a specific job.

    Haiku sees the numbered template paragraphs, resume, and job description,
    then returns which paragraph indices to update and what to write in each.
    Only the text content changes — all DOCX formatting is preserved.
    """
    template_doc = Document(cover_letter_path)
    template_numbered = _numbered_paragraphs(template_doc)

    prompt_template = PROMPT_PATH.read_text()
    substitutions = {
        "template_numbered": template_numbered,
        "resume_text": resume_text,
        "company": job.get("company", ""),
        "job_title": job.get("title", ""),
        "job_description": job.get("description", ""),
    }
    prompt = prompt_template
    for key, value in substitutions.items():
        prompt = prompt.replace(f"{{{key}}}", value)

    message = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    text = message.content[0].text

    # Extract JSON array from response
    json_match = re.search(r"\[.*\]", text, re.DOTALL)
    if not json_match:
        raise ValueError(f"No JSON array found in Claude response:\n{text}")

    replacements = json.loads(json_match.group())

    # Copy base cover letter, apply paragraph replacements, then format
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(cover_letter_path, output_path)

    doc = Document(output_path)
    _apply_paragraph_replacements(doc, replacements)
    _remove_template_remnants(doc)
    _apply_formatting(doc, font_pt=11.0, margin_top=0.5, margin_bottom=0.5)
    doc.save(output_path)

    return output_path
